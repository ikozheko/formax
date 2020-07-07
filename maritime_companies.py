import itertools

from docx import Document
from bs4 import BeautifulSoup
import peewee
from peewee import Model, CharField, ForeignKeyField
from environs import Env
from playhouse.db_url import connect
from tqdm import tqdm
import re
import sys
import os

env = Env()
env.read_env('dev.env')

doc = Document('shipowners_and_shipmanagers.docx')

def _get_text(row):
    data = []
    for paragraph in row.cells[0].paragraphs:
        doc = BeautifulSoup(paragraph._p.xml, 'lxml')            
        words = doc.find_all('w:t')           
        words = filter(lambda item: item.string.strip()!='', words)
        data.append(' '.join([word.string.strip() for word in words if word.string.strip()]))
    return ''.join(data)


def _get_phone(text):
    for i in text:
        if i not in ' -()+,.;/' and not i.isdigit():
            return None

    return text


def _get_email(text):
    if '@' not in text:
        return None
    
    return text


def _get_link(text):
    if not text.startswith('http') and not text.startswith('www'):
        return None

    return text


def parse_table(row_gen):
    obj = {}

    while True:
        row = next(row_gen)
        tr = BeautifulSoup(row._tr.xml, 'lxml')            
        bold_font = tr.find('w:b')  
        if bold_font is None:
            continue

        break

    text = _get_text(row)
    obj['name'] = text

    fields = ['address', 'phone', 'email', 'site', 'description']

    row = next(row_gen)
    text = _get_text(row)
    while fields:

        field = fields.pop(0)

        if field == 'address':
            if (
                _get_phone(text) is None 
                and _get_email(text) is None
                and _get_link(text) is None
            ):
                obj['address'] = text

                row = next(row_gen)
                text = _get_text(row)
                continue

        elif field == 'phone':
            if _get_phone(text) is not None:
                text = re.sub('[\s+\.]', '', text)
                text = re.split('[;,/]', text)
                if isinstance(text, list):
                    text = [phone for phone in text if len(phone)]
                obj['phone'] = text

                row = next(row_gen)
                text = _get_text(row)
                continue

        elif field == 'email':
            if _get_email(text) is not None:
                text = re.sub('\s+', '', text)
                text = re.split(',', text)
                if isinstance(text, list):
                    text = [email for email in text if len(email)]                
                obj['email'] = text

                row = next(row_gen)
                text = _get_text(row)
                continue

        elif field == 'site':
            if _get_link(text) is not None:
                text = re.sub('\s+', '', text)
                obj['site'] = text

                row = next(row_gen)
                text = _get_text(row)
                continue

        elif field == 'description':
            obj['description'] = text

    return obj

db = connect(env('DATABASE_URL'))

class BaseModel(Model):
    class Meta:
        database = db

class DocxCompany(BaseModel):
    name = CharField()
    address = CharField(null=True)
    description = CharField(null=True)
    def __str__(self):
        return self.name

class DocxPhone(BaseModel):
    company = ForeignKeyField(DocxCompany, on_delete='CASCASE')
    number = CharField()
    def __str__(self):
        return self.number

class DocxEmail(BaseModel):
    company = ForeignKeyField(DocxCompany, on_delete='CASCASE')
    address = CharField()
    def __str__(self):
        return self.address

class DocxSite(BaseModel):
    company = ForeignKeyField(DocxCompany, on_delete='CASCASE')
    url = CharField()
    def __str__(self):
        return self.address

# db.create_tables((DocxCompany, DocxPhone, DocxEmail, DocxSite))
def parse_tables(row_gen):
    while True:
        try:
            yield parse_table(row_gen)
        except StopIteration:
            return

row_gen = (
    row
    for table in doc.tables
    for row in table.rows
)

# for tables in tqdm(zip(doc.tables, doc.tables[1:] + [None]), total=len(doc.tables)):

i = 0
for obj in tqdm(parse_tables(row_gen)):
    i += 1
    
    company = DocxCompany.create(
        name=obj['name'],
        address=obj['address'] if obj.get('address') else None,
        description=obj['description'] if obj.get('description') else None,
    )

    if 'phone' in obj.keys():
        if isinstance(obj['phone'], list):
            obj_list = [{'company_id': company.id, 'number': number} for number in obj['phone']]
            DocxPhone.insert_many(rows=obj_list).execute()
        else:
            DocxPhone.create(company_id=company.id, number=obj['phone'])
    
    if 'email' in obj.keys():
        if isinstance(obj['email'], list):
            obj_list = [{'company_id': company.id, 'address': address} for address in obj['email']]
            DocxEmail.insert_many(rows=obj_list).execute()
        else:
            DocxEmail.create(company_id=company.id, address=obj['email'])

    if 'site' in obj.keys():    
        if isinstance(obj['site'], list):
            obj_list = [{'company_id': company.id, 'url': address} for address in obj['site']]
            DocxSite.insert_many(rows=obj_list).execute()
        else:
            DocxSite.create(company_id=company.id, url=obj['site'])

